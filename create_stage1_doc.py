#!/usr/bin/env python3
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

def set_cell_background(cell, color):
    """Set cell background color."""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._element.get_or_add_tcPr().append(shading_elm)

def add_cell_border(cell, **kwargs):
    """Add borders to table cell."""
    tcPr = cell._element.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')

    for edge in ('top', 'left', 'bottom', 'right'):
        edge_data = kwargs.get(edge)
        if edge_data:
            edge_el = OxmlElement(f'w:{edge}')
            edge_el.set(qn('w:val'), 'single')
            edge_el.set(qn('w:sz'), '12')
            edge_el.set(qn('w:space'), '0')
            edge_el.set(qn('w:color'), 'CCCCCC')
            tcBorders.append(edge_el)

    tcPr.append(tcBorders)

doc = Document()

# Set up page margins
sections = doc.sections
for section in sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# Title
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title.add_run('Stage 1')
title_run.font.size = Pt(36)
title_run.font.bold = True
title_run.font.color.rgb = RGBColor(0, 162, 255)

# Subtitle
subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle_run = subtitle.add_run('Color System & Token Overhaul')
subtitle_run.font.size = Pt(24)
subtitle_run.font.bold = True
subtitle_run.font.color.rgb = RGBColor(0, 162, 255)

# Series info
series = doc.add_paragraph()
series.alignment = WD_ALIGN_PARAGRAPH.CENTER
series_run = series.add_run('Weld Landing Page Overhaul Series')
series_run.font.size = Pt(12)
series_run.font.italic = True
series_run.font.color.rgb = RGBColor(84, 110, 122)

doc.add_paragraph()  # Spacing

# Overview
heading1 = doc.add_paragraph()
h1_run = heading1.add_run('Overview')
h1_run.font.size = Pt(16)
h1_run.font.bold = True
h1_run.font.color.rgb = RGBColor(0, 162, 255)

overview_text = doc.add_paragraph(
    'Stage 1 is the foundation of the entire overhaul. It costs the least effort (pure CSS variable changes) '
    'but delivers the highest immediate perception shift. The goal is to replace the current "dark fintech SaaS" '
    'palette with a color system that Roblox developers instantly recognise as native to their world — drawing '
    'from Roblox Studio\'s interface and LuaU\'s syntax highlighting.'
)
overview_text.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

# Why Current Palette Fails
heading2 = doc.add_paragraph()
h2_run = heading2.add_run('Why the Current Palette Fails')
h2_run.font.size = Pt(14)
h2_run.font.bold = True
h2_run.font.color.rgb = RGBColor(0, 162, 255)

failures = [
    'The current background (#0c0e0f) reads as generic startup dark mode — it lacks the characteristic navy-black depth of Roblox Studio\'s own UI',
    'The dominant orange-hot (#ff5a2d) is a SaaS accent color, not a Roblox-native signal',
    'The cream (#fff5f0) text on near-black creates a warm contrast that feels like a lifestyle brand, not a developer tool',
    'There is no visual language that says "this was built by and for Roblox creators"'
]

for failure in failures:
    p = doc.add_paragraph(failure, style='List Bullet')

# New Palette Philosophy
heading3 = doc.add_paragraph()
h3_run = heading3.add_run('The New Palette Philosophy')
h3_run.font.size = Pt(14)
h3_run.font.bold = True
h3_run.font.color.rgb = RGBColor(0, 162, 255)

doc.add_paragraph('The new color system has four layers:')

layers = [
    ('Studio Dark', 'Background stack inspired by Roblox Studio\'s panel system (deep navy-blacks)'),
    ('Studio Blue', 'The primary interactive color, pulled from Studio\'s action button blue (#00a2ff)'),
    ('LuaU Syntax Palette', '5 functional colors pulled directly from LuaU syntax highlighting themes, each assigned a semantic role in the UI'),
    ('Weld Brand Accent', 'Orange/red retained but demoted to CTA energy only, not the primary UI color'),
]

for layer_name, layer_desc in layers:
    p = doc.add_paragraph()
    p_run = p.add_run(layer_name)
    p_run.font.bold = True
    p.add_run(f' — {layer_desc}')

# Token Reference Table
doc.add_paragraph()
heading4 = doc.add_paragraph()
h4_run = heading4.add_run('Detailed Token Reference')
h4_run.font.size = Pt(14)
h4_run.font.bold = True
h4_run.font.color.rgb = RGBColor(0, 162, 255)

# Create table
table = doc.add_table(rows=1, cols=5)
table.style = 'Light Grid Accent 1'

# Header row
hdr_cells = table.rows[0].cells
headers = ['Token Name', 'Old Value', 'New Value', 'Role', 'Where Used']
for i, header_text in enumerate(headers):
    hdr_cells[i].text = header_text
    set_cell_background(hdr_cells[i], '00a2ff')
    for paragraph in hdr_cells[i].paragraphs:
        for run in paragraph.runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)

# Token data
tokens = [
    ('--bg', '#0c0e0f', '#090b14', 'Base background', 'Page body, hero section'),
    ('--bg2', '#080a0b', '#0e111c', 'Panel background', 'Cards, sidebar panels'),
    ('--bg-surface', '(new)', '#131726', 'Raised surface', 'Explorer cards, panels'),
    ('--bg-hover', '(new)', '#1a1f30', 'Interactive hover', 'Button hover, chips'),
    ('--studio-blue', '(new)', '#00a2ff', 'Primary interactive', 'CTA buttons, links'),
    ('--studio-blue-dark', '(new)', '#0077cc', 'Button pressed', 'Pressed/active buttons'),
    ('--luau-keyword', '(new)', '#c792ea', 'LuaU keyword purple', 'Tags, labels, badges'),
    ('--luau-string', '(new)', '#f78c6c', 'LuaU string orange', 'Values, handles'),
    ('--luau-function', '(new)', '#82aaff', 'LuaU function blue', 'Actions, labels'),
    ('--luau-type', '(new)', '#4ec9b0', 'LuaU type teal', 'Verified, success'),
    ('--luau-variable', '(new)', '#c3e88d', 'LuaU variable green', 'Availability, status'),
    ('--luau-comment', '(new)', '#546e7a', 'LuaU comment grey', 'Metadata, timestamps'),
    ('--orange-hot', '#ff5a2d', '#ff5a2d', 'Weld CTA accent', 'Primary CTA only'),
    ('--cream', '#fff5f0', '#e8eaf6', 'Primary text', 'Body text'),
    ('--terminal-dark', '#05070d', '#080b18', 'Deep terminal', 'Boot background'),
]

for token_name, old_val, new_val, role, used in tokens:
    row_cells = table.add_row().cells
    row_cells[0].text = token_name
    row_cells[1].text = old_val
    row_cells[2].text = new_val
    row_cells[3].text = role
    row_cells[4].text = used

    set_cell_background(row_cells[0], '0e111c')
    for cell in row_cells:
        set_cell_background(cell, '0e111c')
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(232, 234, 246)

    row_cells[0].paragraphs[0].runs[0].font.bold = True
    row_cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 162, 255)

# Page break
doc.add_page_break()

# Semantic Color Roles
heading5 = doc.add_paragraph()
h5_run = heading5.add_run('Semantic Color Roles')
h5_run.font.size = Pt(14)
h5_run.font.bold = True
h5_run.font.color.rgb = RGBColor(0, 162, 255)

doc.add_paragraph('Each LuaU syntax color maps to a specific UI function:')

roles = [
    ('Purple (Keyword)', 'Labels, tags, badges — things that "declare" something'),
    ('Orange (String)', 'Values, handles, names — things that ARE something'),
    ('Blue (Function)', 'Actions, navigation, calls-to-action — things you DO'),
    ('Teal (Type)', 'Verification, trust, success — things that are PROVEN'),
    ('Green (Variable)', 'Availability, live state, open — things that are NOW'),
    ('Grey (Comment)', 'Metadata, timestamps, hints — things that explain'),
]

for role_name, role_desc in roles:
    p = doc.add_paragraph()
    p_run = p.add_run(role_name)
    p_run.font.bold = True
    p.add_run(f' — {role_desc}')

# Background Depth System
doc.add_paragraph()
heading6 = doc.add_paragraph()
h6_run = heading6.add_run('Background Depth System')
h6_run.font.size = Pt(14)
h6_run.font.bold = True
h6_run.font.color.rgb = RGBColor(0, 162, 255)

doc.add_paragraph('The new background system uses four distinct depth levels, creating a visual hierarchy that Roblox developers instantly recognise:')

depth_levels = [
    ('--bg (#090b14)', 'The deepest background layer. Used for the page body and hero section background. This is the "canvas" on which everything else sits.'),
    ('--bg2 (#0e111c)', 'One level raised. Used for card backgrounds, sidebar panels, and container elements. Visibly distinct from the base background.'),
    ('--bg-surface (#131726)', 'Two levels raised. Used for role explorer cards, profile panels, and elevated interactive elements.'),
    ('--bg-hover (#1a1f30)', 'The most raised level. Used for button hover states, chip hover effects, and transient interactive feedback.'),
]

for level_name, level_desc in depth_levels:
    p = doc.add_paragraph()
    p_run = p.add_run(level_name)
    p_run.font.bold = True
    p.add_run(f' — {level_desc}')

depth_note = doc.add_paragraph(
    'This four-level depth system mimics the panel-within-panel nesting that Roblox developers work with daily in Studio. '
    'When users see this hierarchy, it signals "this feels like Studio" without any code changes.'
)
depth_note.runs[0].font.italic = True

# CSS Implementation
doc.add_paragraph()
heading7 = doc.add_paragraph()
h7_run = heading7.add_run('CSS Implementation — globals.css')
h7_run.font.size = Pt(14)
h7_run.font.bold = True
h7_run.font.color.rgb = RGBColor(0, 162, 255)

doc.add_paragraph('Replace the :root variable block in globals.css with the following:')

# CSS Before/After table
css_table = doc.add_table(rows=1, cols=2)
css_table.style = 'Light Grid'

hdr = css_table.rows[0].cells
hdr[0].text = 'BEFORE (Current)'
hdr[1].text = 'AFTER (New)'

for cell in hdr:
    set_cell_background(cell, '00a2ff')
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)

# Before content
before_row = css_table.add_row().cells
before_text = (
    ':root {\n'
    '  --bg: #0c0e0f;\n'
    '  --bg2: #080a0b;\n'
    '  --orange-hot: #ff5a2d;\n'
    '  --cream: #fff5f0;\n'
    '  --roblox-blue: #229bd2;\n'
    '}'
)
before_row[0].text = before_text
set_cell_background(before_row[0], '1a1f30')

# After content
after_text = (
    ':root {\n'
    '  --bg: #090b14;\n'
    '  --bg2: #0e111c;\n'
    '  --bg-surface: #131726;\n'
    '  --bg-hover: #1a1f30;\n'
    '  --studio-blue: #00a2ff;\n'
    '  --studio-blue-dark: #0077cc;\n'
    '  --luau-keyword: #c792ea;\n'
    '  --luau-string: #f78c6c;\n'
    '  --luau-function: #82aaff;\n'
    '  --luau-type: #4ec9b0;\n'
    '  --luau-variable: #c3e88d;\n'
    '  --luau-comment: #546e7a;\n'
    '  --orange-hot: #ff5a2d;\n'
    '  --cream: #e8eaf6;\n'
    '  --terminal-dark: #080b18;\n'
    '}'
)
before_row[1].text = after_text
set_cell_background(before_row[1], '1a1f30')

for cell in before_row:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.color.rgb = RGBColor(232, 234, 246)
            run.font.name = 'Consolas'
            run.font.size = Pt(9)

# What Doesn't Change
doc.add_page_break()

heading8 = doc.add_paragraph()
h8_run = heading8.add_run('What Doesn\'t Change')
h8_run.font.size = Pt(14)
h8_run.font.bold = True
h8_run.font.color.rgb = RGBColor(0, 162, 255)

unchanged = [
    'The orange CTA color (#ff5a2d) stays for the primary "CLAIM BETA INVITE" button. It is Weld\'s brand signal and should remain punchy and distinctive.',
    'The scan line overlay and terminal grid aesthetic stay. They are the right visual direction for the landing page; this Stage 1 update shifts them from dominant to complementary.',
    'Font families remain unchanged. Typography refinement is Stage 2; this stage is purely chromatic.'
]

for item in unchanged:
    doc.add_paragraph(item)

# Effort & Impact
doc.add_paragraph()
heading9 = doc.add_paragraph()
h9_run = heading9.add_run('Effort & Impact Assessment')
h9_run.font.size = Pt(14)
h9_run.font.bold = True
h9_run.font.color.rgb = RGBColor(0, 162, 255)

effort_table = doc.add_table(rows=1, cols=2)
effort_table.style = 'Light Grid'

ehdr = effort_table.rows[0].cells
ehdr[0].text = 'Dimension'
ehdr[1].text = 'Details'

for cell in ehdr:
    set_cell_background(cell, '00a2ff')
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)

effort_data = [
    ('Implementation Time', '30–60 minutes. One file: globals.css (:root only). Zero component changes.'),
    ('Testing Scope', 'Visual regression, contrast checker, cross-browser verification'),
    ('Visual Impact', 'Transformative. Page stops reading as generic SaaS instantly.'),
    ('ROI Ratio', 'High. Minimal effort, maximum perception shift.'),
]

for dim, detail in effort_data:
    erow = effort_table.add_row().cells
    erow[0].text = dim
    erow[1].text = detail
    set_cell_background(erow[0], '0e111c')
    set_cell_background(erow[1], '0e111c')

    erow[0].paragraphs[0].runs[0].font.bold = True
    erow[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 162, 255)

    for para in erow[1].paragraphs:
        for run in para.runs:
            run.font.color.rgb = RGBColor(232, 234, 246)

# Stage 2 Dependency
doc.add_paragraph()
heading10 = doc.add_paragraph()
h10_run = heading10.add_run('Stage 2 Dependency')
h10_run.font.size = Pt(14)
h10_run.font.bold = True
h10_run.font.color.rgb = RGBColor(0, 162, 255)

doc.add_paragraph('Stage 2 (Typography & Syntax Highlighting) depends on this palette being set first. The LuaU syntax colors are applied to:')

deps = [
    'Code block syntax highlighting (keywords, strings, functions, types)',
    'Role badge and tag styling',
    'Status indicators and verification badges',
    'Section headers and subsection labels'
]

for dep in deps:
    doc.add_paragraph(dep, style='List Bullet')

# Implementation Checklist
doc.add_paragraph()
heading11 = doc.add_paragraph()
h11_run = heading11.add_run('Implementation Checklist')
h11_run.font.size = Pt(14)
h11_run.font.bold = True
h11_run.font.color.rgb = RGBColor(0, 162, 255)

checklist = [
    'Update :root CSS variables in globals.css with new token values',
    'Test on light and dark OS theme settings to verify contrast ratios',
    'Verify that hero section, card panels, and buttons render with new palette',
    'Take screenshots of landing page at 1440px and 768px viewports',
    'Run WCAG contrast checker on text over new backgrounds (target AA minimum)',
    'Commit changes with message: "Stage 1: Color system & token overhaul"'
]

for item in checklist:
    doc.add_paragraph(item, style='List Bullet')

# Summary
doc.add_paragraph()
heading12 = doc.add_paragraph()
h12_run = heading12.add_run('Summary')
h12_run.font.size = Pt(14)
h12_run.font.bold = True
h12_run.font.color.rgb = RGBColor(0, 162, 255)

summary = doc.add_paragraph(
    'Stage 1 is a high-ROI update. By shifting from a generic SaaS dark palette to a Studio-informed color system '
    'with LuaU syntax semantics, the landing page immediately signals "this was designed by developers, for developers." '
    'The change is pure CSS — no component refactoring, no JavaScript changes, no DOM alterations. The moment the '
    'background darkens to navy and the primary accent shifts to Studio blue, users will feel the difference.'
)
summary.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

final = doc.add_paragraph('Estimated completion: 30-60 minutes. Estimated visual impact: Transformative.')
final.runs[0].font.italic = True
final.runs[0].font.color.rgb = RGBColor(0, 162, 255)

# Save
doc.save('/sessions/keen-friendly-lovelace/mnt/Weld-app/stage1_color_system.docx')
print('Document created successfully: stage1_color_system.docx')
