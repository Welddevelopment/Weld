#!/usr/bin/env python3
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_bg(cell, color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._element.get_or_add_tcPr().append(shading)

doc = Document()
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Stage 1')
r.font.size = Pt(36)
r.font.bold = True
r.font.color.rgb = RGBColor(0, 162, 255)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Color System & Token Overhaul')
r.font.size = Pt(24)
r.font.bold = True
r.font.color.rgb = RGBColor(0, 162, 255)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Weld Landing Page Overhaul Series')
r.font.size = Pt(12)
r.font.italic = True
r.font.color.rgb = RGBColor(84, 110, 122)

doc.add_paragraph()

p = doc.add_paragraph()
r = p.add_run('Overview')
r.font.size = Pt(16)
r.font.bold = True
r.font.color.rgb = RGBColor(0, 162, 255)

p = doc.add_paragraph('Stage 1 is the foundation of the entire overhaul. It costs the least effort (pure CSS variable changes) but delivers the highest immediate perception shift. The goal is to replace the current "dark fintech SaaS" palette with a color system that Roblox developers instantly recognise as native to their world — drawing from Roblox Studio interface and LuaU syntax highlighting.')
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

p = doc.add_paragraph()
r = p.add_run('Why the Current Palette Fails')
r.font.size = Pt(14)
r.font.bold = True
r.font.color.rgb = RGBColor(0, 162, 255)

for item in ['Current background (#0c0e0f) reads generic—lacks navy-black depth', 'Orange-hot (#ff5a2d) is SaaS, not Roblox-native', 'Cream (#fff5f0) text feels lifestyle-like, not dev tool', 'No visual language saying "built by Roblox creators"']:
    doc.add_paragraph(item, style='List Bullet')

p = doc.add_paragraph()
r = p.add_run('The New Palette Philosophy')
r.font.size = Pt(14)
r.font.bold = True
r.font.color.rgb = RGBColor(0, 162, 255)

doc.add_paragraph('Four layers:')
for name, desc in [('Studio Dark', 'Navy-blacks from Studio panel system'), ('Studio Blue', 'Primary color #00a2ff from Studio'), ('LuaU Syntax', '5 colors from syntax highlighting'), ('Weld Accent', 'Orange kept, demoted to CTA only')]:
    p = doc.add_paragraph()
    r = p.add_run(name)
    r.font.bold = True
    p.add_run(f'—{desc}')

doc.add_paragraph()
p = doc.add_paragraph()
r = p.add_run('Detailed Token Reference')
r.font.size = Pt(14)
r.font.bold = True
r.font.color.rgb = RGBColor(0, 162, 255)

table = doc.add_table(rows=1, cols=5)
hdr = table.rows[0].cells
for i, txt in enumerate(['Token', 'Old', 'New', 'Role', 'Used']):
    hdr[i].text = txt
    set_bg(hdr[i], '00a2ff')
    for para in hdr[i].paragraphs:
        for run in para.runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)

tokens = [('--bg','#0c0e0f','#090b14','Base','Page body'),('--bg2','#080a0b','#0e111c','Panel','Cards'),('--bg-surface','new','#131726','Surface','Panels'),('--bg-hover','new','#1a1f30','Hover','Buttons'),('--studio-blue','new','#00a2ff','Primary','CTA'),('--studio-blue-dark','new','#0077cc','Pressed','Buttons'),('--luau-keyword','new','#c792ea','Keyword','Badges'),('--luau-string','new','#f78c6c','String','Values'),('--luau-function','new','#82aaff','Function','Actions'),('--luau-type','new','#4ec9b0','Type','Success'),('--luau-variable','new','#c3e88d','Variable','Status'),('--luau-comment','new','#546e7a','Comment','Meta'),('--orange-hot','#ff5a2d','#ff5a2d','CTA','Primary'),('--cream','#fff5f0','#e8eaf6','Text','Body'),('--terminal-dark','#05070d','#080b18','Deep','Boot')]

for tok, old, new, role, used in tokens:
    row = table.add_row().cells
    row[0].text = tok
    row[1].text = old
    row[2].text = new
    row[3].text = role
    row[4].text = used
    for cell in row:
        set_bg(cell, '0e111c')
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.color.rgb = RGBColor(232, 234, 246)

doc.add_page_break()

p = doc.add_paragraph()
r = p.add_run('Semantic Color Roles')
r.font.size = Pt(14)
r.font.bold = True
r.font.color.rgb = RGBColor(0, 162, 255)

for role, desc in [('Purple','Labels, tags, badges'),('Orange','Values, handles'),('Blue','Actions, CTAs'),('Teal','Verification, success'),('Green','Availability, live'),('Grey','Metadata, timestamps')]:
    p = doc.add_paragraph()
    r = p.add_run(role)
    r.font.bold = True
    p.add_run(f'—{desc}')

doc.add_paragraph()
p = doc.add_paragraph()
r = p.add_run('Background Depth System')
r.font.size = Pt(14)
r.font.bold = True
r.font.color.rgb = RGBColor(0, 162, 255)

for level, desc in [('--bg (#090b14)','Deepest. Body, hero.'),('--bg2 (#0e111c)','Cards, panels.'),('--bg-surface (#131726)','Explorer, profile.'),('--bg-hover (#1a1f30)','Hover, chips.')]:
    p = doc.add_paragraph()
    r = p.add_run(level)
    r.font.bold = True
    p.add_run(f'—{desc}')

p = doc.add_paragraph('Mimics Studio panel nesting. Users see this and think "Studio."')
p.runs[0].font.italic = True

doc.add_paragraph()
p = doc.add_paragraph()
r = p.add_run('CSS Implementation—globals.css')
r.font.size = Pt(14)
r.font.bold = True
r.font.color.rgb = RGBColor(0, 162, 255)

css = doc.add_table(rows=1, cols=2)
ch = css.rows[0].cells
ch[0].text = 'BEFORE'
ch[1].text = 'AFTER'
for c in ch:
    set_bg(c, '00a2ff')
    for pa in c.paragraphs:
        for ru in pa.runs:
            ru.font.bold = True
            ru.font.color.rgb = RGBColor(255, 255, 255)

b = ':root {\n  --bg:#0c0e0f;\n  --bg2:#080a0b;\n}'
a = ':root {\n  --bg:#090b14;\n  --bg2:#0e111c;\n  --studio-blue:#00a2ff;\n  --luau-keyword:#c792ea;\n}'

row = css.add_row().cells
row[0].text = b
row[1].text = a
for c in row:
    set_bg(c, '1a1f30')
    for pa in c.paragraphs:
        for ru in pa.runs:
            ru.font.color.rgb = RGBColor(232, 234, 246)
            ru.font.name = 'Courier'
            ru.font.size = Pt(8)

doc.add_page_break()

p = doc.add_paragraph()
r = p.add_run('What Doesn\'t Change')
r.font.size = Pt(14)
r.font.bold = True
r.font.color.rgb = RGBColor(0, 162, 255)

doc.add_paragraph('Orange CTA (#ff5a2d) stays for "CLAIM BETA INVITE"')
doc.add_paragraph('Scan line and terminal grid stay—complementary')
doc.add_paragraph('Font families unchanged—Stage 2')

doc.add_paragraph()
p = doc.add_paragraph()
r = p.add_run('Effort & Impact')
r.font.size = Pt(14)
r.font.bold = True
r.font.color.rgb = RGBColor(0, 162, 255)

eft = doc.add_table(rows=1, cols=2)
eh = eft.rows[0].cells
eh[0].text = 'Dimension'
eh[1].text = 'Details'
for c in eh:
    set_bg(c, '00a2ff')
    for pa in c.paragraphs:
        for ru in pa.runs:
            ru.font.bold = True
            ru.font.color.rgb = RGBColor(255, 255, 255)

for d, de in [('Time','30-60 min. One file.'),('Testing','Visual, contrast, browser'),('Impact','Transformative'),('ROI','High')]:
    r = eft.add_row().cells
    r[0].text = d
    r[1].text = de
    set_bg(r[0], '0e111c')
    set_bg(r[1], '0e111c')
    for pa in r[1].paragraphs:
        for ru in pa.runs:
            ru.font.color.rgb = RGBColor(232, 234, 246)

doc.add_paragraph()
p = doc.add_paragraph()
r = p.add_run('Stage 2 Dependency')
r.font.size = Pt(14)
r.font.bold = True
r.font.color.rgb = RGBColor(0, 162, 255)

doc.add_paragraph('Stage 2 depends on palette. Colors apply to syntax, badges, status, headers.')

doc.add_paragraph()
p = doc.add_paragraph()
r = p.add_run('Implementation Checklist')
r.font.size = Pt(14)
r.font.bold = True
r.font.color.rgb = RGBColor(0, 162, 255)

for item in ['Update :root in globals.css','Test light/dark themes','Verify hero/cards/buttons','Screenshot 1440px, 768px','WCAG contrast check','Commit']:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph()
p = doc.add_paragraph()
r = p.add_run('Summary')
r.font.size = Pt(14)
r.font.bold = True
r.font.color.rgb = RGBColor(0, 162, 255)

p = doc.add_paragraph('Stage 1 is high-ROI. Pure CSS. No refactoring, JS, DOM changes. Signals "designed for developers."')
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

p = doc.add_paragraph('Estimated: 30-60 minutes. Visual impact: Transformative.')
p.runs[0].font.italic = True
p.runs[0].font.color.rgb = RGBColor(0, 162, 255)

doc.save('/sessions/keen-friendly-lovelace/mnt/Weld-app/stage1_color_system.docx')
print('Created: stage1_color_system.docx')
