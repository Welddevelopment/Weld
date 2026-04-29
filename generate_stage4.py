#!/usr/bin/env python3
"""Generate Stage 4 Word document for Weld landing page overhaul."""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Create document
doc = Document()

# Set margins
sections = doc.sections
for section in sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# Define colors
DARK_BLUE = RGBColor(31, 71, 136)
MEDIUM_BLUE = RGBColor(46, 92, 138)
LIGHT_BLUE = RGBColor(62, 124, 184)
DARK_GRAY = RGBColor(102, 102, 102)
LIGHT_GRAY = RGBColor(153, 153, 153)

# TITLE PAGE
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("Stage 4")
run.font.size = Pt(32)
run.font.bold = True
run.font.color.rgb = DARK_BLUE

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("Section-by-Section Component Overhaul")
run.font.size = Pt(28)
run.font.bold = True
run.font.color.rgb = DARK_BLUE
subtitle.space_before = Pt(12)

series = doc.add_paragraph()
series.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = series.add_run("Weld Landing Page Overhaul Series")
run.font.size = Pt(24)
run.font.italic = True
run.font.color.rgb = DARK_GRAY
series.space_before = Pt(24)

spec = doc.add_paragraph()
spec.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = spec.add_run("Technical Specification & Implementation Guide")
run.font.size = Pt(22)
run.font.color.rgb = DARK_GRAY
spec.space_before = Pt(12)

date = doc.add_paragraph()
date.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = date.add_run("April 2026")
run.font.size = Pt(20)
run.font.color.rgb = LIGHT_GRAY
date.space_before = Pt(12)

# PAGE BREAK
doc.add_page_break()

# OVERVIEW
heading = doc.add_heading("Overview", level=1)
heading.runs[0].font.color.rgb = DARK_BLUE

doc.add_paragraph("Stage 4 is the most substantial visual redesign in the Weld landing page overhaul series. Where Stages 1-3 changed system-level aesthetics (colors, typography, and motion), Stage 4 performs a deep visual redesign of each individual section using a specific Roblox-native UI metaphor.")
doc.add_paragraph("Every section of the landing page is mapped to a recognisable element of the Roblox ecosystem: Studio, Discord, the Explorer, the Properties panel. This design approach makes the page feel like it was built from inside the tools that Roblox developers already live in daily, creating immediate familiarity and trust.")
doc.add_paragraph("Each section redesign is a contained component change that can be shipped independently. The recommended implementation order prioritizes sections with the highest visual impact and developer resonance.")

# METAPHOR MAP
heading = doc.add_heading("The Metaphor Map", level=2)
heading.runs[0].font.color.rgb = MEDIUM_BLUE

doc.add_paragraph("Each landing page section is mapped to a Roblox UI element. The table below shows the current look, new metaphor, and inspiration source for every component redesign.")

# Create table
table = doc.add_table(rows=11, cols=4)
table.style = 'Light Grid Accent 1'

# Set column widths
widths = (Inches(1.4), Inches(1.7), Inches(1.7), Inches(1.8))
for row in table.rows:
    for idx, width in enumerate(widths):
        row.cells[idx].width = width

# Header row
header_cells = table.rows[0].cells
headers = ["Section", "Current Look", "New Metaphor", "Inspiration Source"]
for idx, header_text in enumerate(headers):
    cell = header_cells[idx]
    cell.text = header_text
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), '1F4788')
    cell._element.get_or_add_tcPr().append(shading_elm)

# Data rows
data = [
    ["Navigation Header", "Dark pill bar with logo and CTA", "Roblox Studio Toolbar", "Studio top toolbar with icon groups"],
    ["Hero Left Panel", "Terminal boot sequence", "Studio Output Window", "Studio Output panel"],
    ["Hero Right Panel", "Rotating profile card", "Studio Properties Inspector", "Studio Properties panel"],
    ["Discord Chaos Section", "Styled div messages", "Real Discord server UI", "discord.com dark theme"],
    ["Role Explorer", "Horizontal chip pills", "Studio Explorer Tree", "Studio Explorer panel"],
    ["How It Works", "Terminal command blocks", "LuaU Script Editor Tabs", "Studio Script tab editor"],
    ["Stats Section", "Plain number grid", "Roblox game server stats", "Roblox game details page"],
    ["Referral Rewards", "Basic reward cards", "Roblox Badge unlock UI", "Roblox badge award modal"],
    ["CTA Section", "Dark email form panel", "Roblox Studio Welcome Screen", "Studio New Project welcome card"],
    ["FAQ", "Simple accordion", "Studio Properties collapsible", "Studio property group collapse"]
]

for row_idx, row_data in enumerate(data):
    row = table.rows[row_idx + 1]
    for col_idx, cell_text in enumerate(row_data):
        cell = row.cells[col_idx]
        cell.text = cell_text
        bg_color = 'F5F5F5' if row_idx % 2 == 0 else 'FFFFFF'
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), bg_color)
        cell._element.get_or_add_tcPr().append(shading_elm)

# PAGE BREAK
doc.add_page_break()

# SECTION REDESIGNS
heading = doc.add_heading("Section Redesigns", level=2)
heading.runs[0].font.color.rgb = MEDIUM_BLUE

doc.add_paragraph("The following sections detail the visual redesign for each component on the landing page. Each redesign provides specific implementation guidance including styling directives, element structure, and visual references.")

# NAVIGATION HEADER
heading = doc.add_heading("Navigation Header: Studio Toolbar", level=3)
heading.runs[0].font.color.rgb = LIGHT_BLUE

p = doc.add_paragraph()
r = p.add_run("Current: ")
r.bold = True
p.add_run("A rounded pill bar (header-rail class) with the Weld logo, audience toggle in the centre, and a CTA button on the right.")

p = doc.add_paragraph()
r = p.add_run("New Studio Toolbar Design: ")
r.bold = True
p.add_run("Remove the rounded pill. The toolbar now spans full width with a flat bottom border, exactly like Studio's toolbar.")

doc.add_paragraph("Weld logo mark + wordmark (unchanged) | thin vertical divider | MODE toggle as icon+text tab strip", style='List Bullet')
doc.add_paragraph("Search/filter bar styled like Studio's search box (dark inset, magnifier icon)", style='List Bullet')
doc.add_paragraph("Notification bell icon | avatar circle | primary CTA button in Studio blue", style='List Bullet')
doc.add_paragraph("1px border in rgba(255,255,255,0.08), no shadow", style='List Bullet')

# HERO LEFT PANEL
heading = doc.add_heading("Hero Left Panel: Studio Output Window", level=3)
heading.runs[0].font.color.rgb = LIGHT_BLUE

p = doc.add_paragraph()
r = p.add_run("Current: ")
r.bold = True
p.add_run("Terminal-rail div with scan overlay and boot lines.")

p = doc.add_paragraph()
r = p.add_run("New Output Window Design: ")
r.bold = True
p.add_run("A fake Studio panel chrome with top bar labeled 'Output', X and icons (right, decorative).")

doc.add_paragraph("Boot sequence lines formatted as Studio Output", style='List Bullet')
doc.add_paragraph("Filter bar styled like Studio's Filter Output--pills: ALL | PRINT | WARN | ERROR", style='List Bullet')
doc.add_paragraph("Dark background (#0d1117) distinct from page background", style='List Bullet')
doc.add_paragraph("2px colored border matching current audience mode color", style='List Bullet')

# HERO RIGHT PANEL
heading = doc.add_heading("Hero Right Panel: Studio Properties Inspector", level=3)
heading.runs[0].font.color.rgb = LIGHT_BLUE

p = doc.add_paragraph()
r = p.add_run("Current: ")
r.bold = True
p.add_run("Floating profile card with avatar, handle, stats in a card layout.")

p = doc.add_paragraph()
r = p.add_run("New Properties Inspector Design: ")
r.bold = True
p.add_run("Panel chrome header reads 'Properties - DeveloperProfile'.")

doc.add_paragraph("Two-column table: Left (label) and right (value). Properties: DisplayName, Role, Rate, Availability, ShippedGames, TotalVisits, Timezone, ResponseTime", style='List Bullet')
doc.add_paragraph("Alternating subtle row shading (#131726 / #0e111c)", style='List Bullet')
doc.add_paragraph("Monospace font values; verified properties have small teal checkmark", style='List Bullet')
doc.add_paragraph("Avatar sits above panel in small 'instance icon' style", style='List Bullet')

# DISCORD CHAOS SECTION
heading = doc.add_heading("Discord Chaos Section: Real Discord Dark Mode", level=3)
heading.runs[0].font.color.rgb = LIGHT_BLUE

p = doc.add_paragraph()
r = p.add_run("Current: ")
r.bold = True
p.add_run("Styled divs that look 'Discord-inspired' but don't pass for the real thing.")

p = doc.add_paragraph()
r = p.add_run("New: ")
r.bold = True
p.add_run("Actually looks like Discord's dark mode.")

doc.add_paragraph("Panel background: Discord's exact dark: #313338", style='List Bullet')
doc.add_paragraph("Message layout: Avatar circle (40px) + username + timestamp + message body", style='List Bullet')
doc.add_paragraph("Username colors: Discord's role color system--hue-based colours", style='List Bullet')
doc.add_paragraph("Sidebar partially visible on left, showing '#dev-hiring' channel", style='List Bullet')
doc.add_paragraph("Spam/noise: Discord's 'This message was blocked' treatment", style='List Bullet')
doc.add_paragraph("Deleted messages: Show '[Original Message Deleted]' in Discord's grey italic", style='List Bullet')
doc.add_paragraph("Transition: Discord panel left, Weld output right, glowing compile line", style='List Bullet')

# PAGE BREAK
doc.add_page_break()

# ROLE EXPLORER
heading = doc.add_heading("Role Explorer: Studio Explorer Tree", level=3)
heading.runs[0].font.color.rgb = LIGHT_BLUE

p = doc.add_paragraph()
r = p.add_run("Current: ")
r.bold = True
p.add_run("Horizontal pill chips with a card showing role stats.")

p = doc.add_paragraph()
r = p.add_run("New Studio Explorer Design: ")
r.bold = True
p.add_run("Full Explorer panel mockup: dark panel, 'Explorer' header with search box.")

doc.add_paragraph("Hierarchy: Workspace > Roles > [role folders]", style='List Bullet')
doc.add_paragraph("Role nodes: Folder nodes (closed/open). Children: Rate, Availability, Skills, ShippedWork with icons", style='List Bullet')
doc.add_paragraph("Icons: Studio Explorer's exact icon set style (16px icons before each node)", style='List Bullet')
doc.add_paragraph("Selection highlight: Active/selected role has full-width Studio blue", style='List Bullet')
doc.add_paragraph("Scrollbar: Dark scrollbar styled like Studio's", style='List Bullet')

# HOW IT WORKS
heading = doc.add_heading("How It Works: LuaU Script Editor Tabs", level=3)
heading.runs[0].font.color.rgb = LIGHT_BLUE

p = doc.add_paragraph()
r = p.add_run("Current: ")
r.bold = True
p.add_run("Three command steps with terminal lines.")

p = doc.add_paragraph()
r = p.add_run("New Script Editor Design: ")
r.bold = True
p.add_run("Full script editor mock: tab bar with three script tabs.")

doc.add_paragraph("Tabs: 'weld_auth.luau', 'weld_profile.luau', 'weld_discover.luau'", style='List Bullet')
doc.add_paragraph("Active tab: White text, bottom border in studio blue", style='List Bullet')
doc.add_paragraph("Line numbers: Left gutter (1-15 lines)", style='List Bullet')
doc.add_paragraph("Syntax highlighting: Full LuaU highlighting using Stage 1 palette", style='List Bullet')
doc.add_paragraph("Run button: Below editor; Studio's green play button aesthetic", style='List Bullet')
doc.add_paragraph("Completion: Appears as comment -- PROOF_VERIFIED", style='List Bullet')

# FINAL CTA SECTION
heading = doc.add_heading("Final CTA Section: Studio Welcome Screen", level=3)
heading.runs[0].font.color.rgb = LIGHT_BLUE

p = doc.add_paragraph()
r = p.add_run("Current: ")
r.bold = True
p.add_run("Dark panel with email input.")

p = doc.add_paragraph()
r = p.add_run("New Welcome Screen Design: ")
r.bold = True
p.add_run("Mimics Roblox Studio welcome: centred card with Studio blue gradient top.")

doc.add_paragraph("Header: 'Start a New Project' -> 'Start Your Weld Profile'", style='List Bullet')
doc.add_paragraph("Developer path tiles: SCRIPTER, BUILDER, DESIGNER (like Studio's Baseplate/Terrain)", style='List Bullet')
doc.add_paragraph("Email form: Below tiles; styled like Studio's 'Sign in to Roblox Studio' input", style='List Bullet')
doc.add_paragraph("Micro copy: 'Your shipped work speaks before you do.'", style='List Bullet')

# IMPLEMENTATION & EFFORT
heading = doc.add_heading("Implementation Order & Effort", level=2)
heading.runs[0].font.color.rgb = MEDIUM_BLUE

doc.add_paragraph("Each section redesign is a contained component change that can be shipped independently. The following order prioritizes sections with the highest visual impact and developer resonance.")

heading = doc.add_heading("Recommended Sequence:", level=3)
heading.runs[0].font.color.rgb = LIGHT_BLUE

doc.add_paragraph("Discord Chaos section--Most differentiation; currently furthest from real Discord", style='List Bullet')
doc.add_paragraph("Hero right panel--Immediately readable to any Roblox developer", style='List Bullet')
doc.add_paragraph("Role Explorer--Most interactive; highest engagement section", style='List Bullet')
doc.add_paragraph("How It Works--Pairs with Stage 2's LuaU syntax work", style='List Bullet')
doc.add_paragraph("Navigation header--Last, visible on all sections", style='List Bullet')

heading = doc.add_heading("Estimated Effort & Impact:", level=3)
heading.runs[0].font.color.rgb = LIGHT_BLUE

doc.add_paragraph("Total time: 15-25 hours across all sections", style='List Bullet')
doc.add_paragraph("Highest impact: Discord Chaos and Hero Properties Inspector alone transform impression", style='List Bullet')
doc.add_paragraph("Per-section: 2-4 hours depending on complexity", style='List Bullet')

# CLOSING
p = doc.add_paragraph()
r = p.add_run("Stage 4 transforms the landing page from generic SaaS to a deeply contextual, tool-native experience. By mapping each section to a recognisable Roblox UI metaphor, we make the platform immediately legible to developers who live in Studio daily.")
r.italic = True

p = doc.add_paragraph()
r = p.add_run("Execution begins with maximum-impact sections (Discord, Properties) and expands to complete the visual overhaul. Each section ships independently, enabling rapid iteration and feedback.")
r.italic = True

# Save document
output_path = "/sessions/keen-friendly-lovelace/mnt/Weld-app/stage4_components.docx"
doc.save(output_path)
print("SUCCESS: Document saved to " + output_path)
