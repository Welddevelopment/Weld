#!/usr/bin/env python3
"""Generate Stage 2 DOCX document using python-docx"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_border(paragraph, **kwargs):
    """Add borders to a paragraph"""
    p = paragraph._element
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    for border_name in ['top', 'left', 'bottom', 'right']:
        border_el = OxmlElement(f'w:{border_name}')
        border_el.set(qn('w:val'), 'single')
        border_el.set(qn('w:sz'), '12')
        border_el.set(qn('w:space'), '0')
        border_el.set(qn('w:color'), 'CCCCCC')
        pBdr.append(border_el)
    pPr.append(pBdr)

doc = Document()

# Set up styles
style = doc.styles['Normal']
style.font.name = 'Arial'
style.font.size = Pt(11)

# Title
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title.add_run("Stage 2 — LuaU Code Aesthetic & Typography")
title_run.font.size = Pt(18)
title_run.font.bold = True
title_run.font.color.rgb = RGBColor(31, 73, 125)

# Subtitle
subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle_run = subtitle.add_run("Weld Landing Page Overhaul Series")
subtitle_run.font.size = Pt(12)
subtitle_run.font.color.rgb = RGBColor(68, 84, 106)
subtitle_run.italic = True

# Overview heading
h2 = doc.add_heading("Overview", level=2)
h2.style.font.color.rgb = RGBColor(46, 92, 138)

# Overview paragraph
overview = doc.add_paragraph(
    "Stage 2 transforms how text and code look on the page. The current boot sequence lines read like a generic CLI tool. "
    "The \"How It Works\" command steps look like terminal output. Neither signals \"Roblox Studio\" to a Roblox developer. "
    "This stage makes those elements look exactly like the tools developers already live inside: the Roblox Studio Output window, "
    "the Script editor, and the Explorer tree panel."
)

# Problem heading
h2 = doc.add_heading("The Problem with Current Typography", level=2)

# Problem bullets
problems = [
    "Boot sequence lines are just monospace text with letter-spacing—they could be from any dark SaaS tool",
    "The \"How It Works\" command blocks look like bash terminal output, not Luau code",
    "There is no visual language in the type system that says \"this is a Roblox creator tool\"",
    "Information density is too low—Roblox Studio UI is compact and dense; the current page feels airy and startup-y",
    "The shimmer animation on every CTA button cheapens the aesthetic—it's a Webflow template trick"
]
for problem in problems:
    doc.add_paragraph(problem, style='List Bullet')

# Change 1
h2 = doc.add_heading("Change 1: The LuaU Code Block Component", level=2)
doc.add_paragraph(
    "The three command steps in \"How It Works\" currently show terminal-style strings. "
    "These must be replaced with a proper LuaU code block component that syntax-highlights like a real Luau script editor."
)

h3 = doc.add_heading("Component Specifications", level=3)
specs = [
    "Background: #0d1117 (darker than page bg, like the script editor tab background)",
    "Font: JetBrains Mono or Fira Code (monospace, same family as Studio's script editor)",
    "Tab bar at the top with fake script names: weld_auth.luau, weld_profile.luau, weld_discover.luau",
    "Active tab has a bottom border in --studio-blue",
    "Line numbers in the left gutter (grey, like Studio)",
    "Syntax coloring using the Stage 1 LuaU palette"
]
for spec in specs:
    doc.add_paragraph(spec, style='List Bullet')

# Syntax table heading
h3 = doc.add_heading("Syntax Color Palette", level=3)

# Create syntax table
table = doc.add_table(rows=7, cols=3)
table.style = 'Light Grid Accent 1'
header_cells = table.rows[0].cells
header_cells[0].text = "Token Type"
header_cells[1].text = "Color"
header_cells[2].text = "Examples"

data = [
    ["Keywords", "#c792ea", "local, function, end, return, if, then"],
    ["Strings", "#f78c6c", "\"Scripter\", \"45 R$/hr\""],
    ["Function Names", "#82aaff", "WeldProfile.build, verify, discover"],
    ["Numbers", "#f07178", "45, 17.3"],
    ["Comments", "#546e7a", "-- PROOF_VERIFIED"],
    ["Variables", "#c3e88d", "dev, profile, result"]
]

for i, row_data in enumerate(data, 1):
    row = table.rows[i].cells
    row[0].text = row_data[0]
    row[1].text = row_data[1]
    row[2].text = row_data[2]

# Example code
h3 = doc.add_heading("Example LuaU Code", level=3)
code_block = [
    "-- weld_auth.luau",
    "local WeldProfile = {}",
    "",
    "function WeldProfile.build(dev)",
    "  dev.role = \"Scripter\"",
    "  dev.rate = \"45 R$/hr\"",
    "  dev.availability = \"Open now\"",
    "  return WeldProfile.verify(dev)",
    "end -- PROOF_VERIFIED"
]
for line in code_block:
    p = doc.add_paragraph(line)
    p.paragraph_format.left_indent = Inches(0.5)
    for run in p.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)

# Change 2
h2 = doc.add_heading("Change 2: The Studio Output Panel", level=2)
doc.add_paragraph(
    "The hero boot sequence currently looks generic. It should look like the Roblox Studio Output window. Key changes:"
)

output_changes = [
    "Add a fake \"Output\" panel header with the Studio-style title bar (dark bar, \"Output\" text left, close X right)",
    "Prefix each line with a bracketed timestamp: [12:04:01.342]",
    "Color-code by severity using Studio's Output colors",
    "The final \"READY_FOR_DISCOVERY\" line should be bold teal, like a print() success in Studio"
]
for change in output_changes:
    doc.add_paragraph(change, style='List Bullet')

# Output colors table
h3 = doc.add_heading("Output Colors by Severity", level=3)
table = doc.add_table(rows=5, cols=3)
table.style = 'Light Grid Accent 1'
header_cells = table.rows[0].cells
header_cells[0].text = "Severity"
header_cells[1].text = "Color"
header_cells[2].text = "Usage"

severity_data = [
    ["Info", "White/Light grey", "Status messages"],
    ["Success", "#4ec9b0 (Teal)", "Completion messages"],
    ["Warning", "#ffb347 (Amber)", "Non-fatal alerts"],
    ["Error", "#f07178 (Red)", "Error messages"]
]

for i, row_data in enumerate(severity_data, 1):
    row = table.rows[i].cells
    row[0].text = row_data[0]
    row[1].text = row_data[1]
    row[2].text = row_data[2]

# Boot sequence
h3 = doc.add_heading("New Boot Sequence Format", level=3)
boot_lines = [
    "[12:04:01.001]  Booting weld.roster v2.0...",
    "[12:04:01.120]  Loading developer lane...",
    "[12:04:01.245]  Scanning shipped work... OK",
    "[12:04:01.471]  Verified: 17.3M total visits",
    "[12:04:01.620]  Matching studio filters...",
    "[12:04:01.781]  READY_FOR_DISCOVERY"
]
for line in boot_lines:
    p = doc.add_paragraph(line)
    p.paragraph_format.left_indent = Inches(0.5)
    for run in p.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)

# Change 3
h2 = doc.add_heading("Change 3: Explorer Tree Typography for Role Explorer", level=2)
doc.add_paragraph(
    "The current role chips (SCRIPTER, UI/UX, VFX, etc.) are flat pills in a row. "
    "They should be restructured to look like the Roblox Studio Explorer tree panel."
)

h3 = doc.add_heading("Tree Panel Design", level=3)
tree_design = [
    "Font: monospace, 11px, tighter line-height",
    "Each role becomes a \"folder\" node with a arrow indicator",
    "Clicking expands to show child nodes: role stats, skills as child items",
    "The active/selected node has a Studio-blue highlight row (full width, like Explorer selection)"
]
for design in tree_design:
    doc.add_paragraph(design, style='List Bullet')

# Tree structure
h3 = doc.add_heading("Tree Structure", level=3)
tree_lines = [
    " Workspace",
    "  Scripter          [open now  45 R$/hr]",
    "    LUAU",
    "    OOP",
    "    DATASTORESERVICE",
    "  UI / UX           [open this week  35 R$/hr]",
    "  VFX               [3 slots open  30 R$/hr]",
    "  Builder           [sprint lane  25 R$/hr]"
]
for line in tree_lines:
    p = doc.add_paragraph(line)
    p.paragraph_format.left_indent = Inches(0.5)
    for run in p.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)

# Change 4
h2 = doc.add_heading("Change 4: Information Density & Font Sizing", level=2)
doc.add_paragraph(
    "Current font sizes are too generous. The page breathes too much for a tool aimed at developers. Proposed changes:"
)

density_changes = [
    "Body text: reduce from implied 14–15px to 13px",
    "Mono text in boot sequence: reduce from 11px to 10px (tighter, denser)",
    "Add a \"compact mode\" variant for the stats section—pack numbers closer together",
    "Replace the loose section padding with tighter values in key areas",
    "The hero profile card metadata should pack more rows in the same space"
]
for change in density_changes:
    doc.add_paragraph(change, style='List Bullet')

# Change 5
h2 = doc.add_heading("Change 5: Kill the Shimmer, Add Targeted Glow", level=2)
doc.add_paragraph(
    "The current .command-button::before has a shimmer animation running at all times. "
    "This is a 2019 Webflow template move. Replace with:"
)

shimmer_changes = [
    "No shimmer on idle state",
    "On hover only: a clean glow box-shadow using --studio-blue",
    "CTA button hover: box-shadow: 0 0 32px rgba(0, 162, 255, 0.45)",
    "Border color shifts to studio blue on hover instead of just brightening"
]
for change in shimmer_changes:
    doc.add_paragraph(change, style='List Bullet')

# Implementation
h2 = doc.add_heading("Implementation Order", level=2)
impl_steps = [
    "Create a SyntaxBlock component in src/components/SyntaxBlock.tsx",
    "Update the HOW_COMMANDS rendering in MarketingPage.tsx to use SyntaxBlock",
    "Update the boot sequence container to add the Output panel header",
    "Update .terminal-chip styles to Explorer tree layout",
    "Update globals.css to remove shimmer from idle state",
    "Add targeted glow CSS for hover states"
]
for step in impl_steps:
    doc.add_paragraph(step, style='List Bullet')

# Effort section
h2 = doc.add_heading("Effort & Impact", level=2)
doc.add_paragraph(
    "Estimated implementation time: 3–5 hours. Files changed: globals.css, MarketingPage.tsx (2 sections), "
    "new SyntaxBlock component. This stage is where Roblox developers first feel \"they get us.\""
)

h3 = doc.add_heading("Key Metrics", level=3)
table = doc.add_table(rows=5, cols=2)
table.style = 'Light Grid Accent 1'
header_cells = table.rows[0].cells
header_cells[0].text = "Metric"
header_cells[1].text = "Value"

metrics_data = [
    ["Estimated Implementation Time", "3–5 hours"],
    ["Files Modified", "3 (globals.css, MarketingPage.tsx, SyntaxBlock.tsx)"],
    ["Components Created", "1 (SyntaxBlock)"],
    ["Impact", "Developer alignment—signals \"this tool speaks our language\""]
]

for i, row_data in enumerate(metrics_data, 1):
    row = table.rows[i].cells
    row[0].text = row_data[0]
    row[1].text = row_data[1]

# Summary
h2 = doc.add_heading("Summary", level=2)
doc.add_paragraph(
    "Stage 2 moves Weld beyond a generic dark SaaS aesthetic into a cohesive visual system rooted in Roblox's own development tools. "
    "By adopting the typography, code styling, and UI patterns that Roblox developers see every day in Studio, we communicate that "
    "Weld is built \"for us, by us.\" The changes are surgical but high-impact: three new components, two CSS updates, and tighter "
    "information density transform the page from aspirational startup pitch to credible developer platform."
)

# Save document
output_path = "C:\\Users\\cubit\\Downloads\\Weld-app\\stage2_typography.docx"
doc.save(output_path)
print(f"Document created successfully: {output_path}")
